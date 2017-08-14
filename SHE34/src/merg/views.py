from django.shortcuts import render,get_object_or_404,redirect,render_to_response
from django.contrib import messages
from django.http import HttpResponse , HttpResponseRedirect
from django.core.urlresolvers import reverse
from django.views.generic import DetailView,CreateView , UpdateView
from HE3.models import Project, Evaluation, Screenshots, Environment
from django.utils import timezone
from docx import Document
import csv
import merg.recommend as rec
from django.contrib.auth.decorators import login_required
import json
from django.views.decorators.csrf import csrf_exempt
from django.core.urlresolvers import reverse_lazy
from HE3.forms import MergeEvaluationForm

def mergeDesktop(request, project_id):

    project = Project.objects.get(pk = project_id);
    if (project.manager == request.user):
        template_name = 'merge/merge.html'
        allEvaluations = project.evaluation_for_project.all()
        alreadyMerged = allEvaluations.filter(merged =True).values_list('mergedFromEvaluations' , flat = True)
        remainedEvals = allEvaluations.filter(merged=False).exclude( pk__in = alreadyMerged)
        # print('all , alreadymerged , remained' , allEvaluations.count() , alreadyMerged.count() , remainedEvals.count())
        if remainedEvals.count() == 0:
            messages.error(request,'<h3>There is no evaluation remained to be merged!</h3>')

        context = {'project': project,
                   'now': timezone.now().date(),
                   'evals': remainedEvals,
                   'mergedEvals': allEvaluations.filter(merged=True),
                   'evaluators': project.evaluators.filter(pk__in = remainedEvals.values_list('evaluator', flat= True).distinct()),
                   }
        return render(request,template_name,context)
    else:
        return HttpResponse('Only Manager can access this page!')

@login_required
def recommend(request , eval_id):

    eval = Evaluation.objects.get(id = eval_id)
    project =eval.ofProject
    resultplace = rec.placebase(eval)
    # resultdes = rec.descriptionBase(eval)
    result = rec.nearestN(eval)
    # context = {'result': eval.ofProject}

    recPlaceBase = project.evaluation_for_project.filter(pk__in =resultplace)
    recDesBase = project.evaluation_for_project.filter(pk__in =result['resultDes'])
    recRecBase = project.evaluation_for_project.filter(pk__in =result['resultRec'])
    recTagseBase = project.evaluation_for_project.filter(pk__in =result['resultTags'])

    if request.is_ajax():
        # template_name = 'HE3/evaluations/e-panel-list.html'
        template_name = 'merge/user-based-evals.html'
        allRecommendedEvals = (recDesBase | recRecBase | recTagseBase).distinct()
        allEvaluations = project.evaluation_for_project.all().exclude(pk=eval_id)
        alreadyMerged = allEvaluations.filter(merged=True).values_list('mergedFromEvaluations', flat=True)

        context = {'evals' : allRecommendedEvals.exclude(pk__in= alreadyMerged) ,
                   'used_evaluations': allRecommendedEvals.filter(pk__in =alreadyMerged),
                   'evaluators': project.evaluators.filter(pk__in = allRecommendedEvals.values_list('evaluator', flat= True).distinct()),
                   }

    else:
        template_name = 'merge/recommendations.html'
        context = {'placeBase': recPlaceBase,
                   'desBase': recDesBase,
                   'recBase': recRecBase,
                   'tagBase': recTagseBase}

    return render(request,template_name=template_name,context= context )

# place base
def recommendPlaceBase(request, eval_id):

    eval = Evaluation.objects.get(id=eval_id)
    project = eval.ofProject
    resultplace = rec.placebase(eval)
    result = {}
    result['evaluations'] = resultplace
    recPlaceBase = project.evaluation_for_project.filter(pk__in =resultplace)

    allEvaluations = project.evaluation_for_project.all().exclude(pk=eval_id)
    alreadyMerged = allEvaluations.filter(merged=True).values_list('mergedFromEvaluations', flat=True)

    if eval.link:
        sameLink = allEvaluations.filter(link__exact=eval.link)
        recPlaceBase = (recPlaceBase | sameLink).distinct()
    return render_to_response(template_name='merge/user-based-evals.html' , context= {'evals' : recPlaceBase.exclude(pk__in=alreadyMerged),
                                                                                      'used_evaluations':recPlaceBase.filter(pk__in=alreadyMerged),
                                                                                      'evaluators' : project.evaluators.filter(pk__in = recPlaceBase.values_list('evaluator', flat= True).distinct()) })

def mergeFields(evalList):
    # result = {k : "" for k in evalList[0].__dict__.keys()}
    result ={}
    stringList = ['description' , 'recommendation' ]
    stringWithKomma = ['place' , 'title' ]
    forAvg = ['frequency' , 'severity']
    for field in stringList :
        allvalues = set([ i.__dict__[field] for i in evalList])
        allvalues = [str(j) + ')  ' + v for j , v in zip(list(range(1 , len(allvalues)+1)) , allvalues) ]
        result [field] = "\n" .join(allvalues)
    for field in stringWithKomma :
        allvalues = set([i.__dict__[field]for i in evalList])
        result[field] = ", ".join(allvalues)
    for field in forAvg :
        allvalues = [int(i.__dict__[field]) for i in evalList]
        if (len(allvalues) > 0):
            result[field] = str(sum(allvalues) / len(allvalues))
        else:
            result[field] = str(1)


    heurPrincips =[]
    for e in evalList:
        heurPrincips.extend(list(e.heurPrincip.all()))
    heurPrincips = list(set(heurPrincips))

    return result , heurPrincips

def mergeScreenshots(resultEval, listEvals):
    for eval in listEvals:
        if eval.merged:
                if eval.mergedScreenshots:
                    resultEval.mergedScreenshots.add(*eval.mergedScreenshots.all())
        elif eval.screenshot:
            screenobject =Screenshots( screenshot =eval.screenshot)
            screenobject.save()
            resultEval.mergedScreenshots.add(screenobject)

# function to merge multiple evaluation
def mergeEvaluations(request , project_id):

    project = Project.objects.get(pk=project_id)
    resultEval = Evaluation(ofProject =project , evaluator=project.manager , merged= True)

    if request.is_ajax():
        ids = request.POST.getlist('ids[]')
        # name = request.POST.get('name')

        if ids and len(ids) > 1:
            evals = Evaluation.objects.filter(pk__in=ids)
            fields , heurPrincips = mergeFields(evals)
            resultEval.__dict__.update(fields)
            # if name:
            #     resultEval.title = name
            # else:
            #     resultEval.title = 'Merged Evaluations'
            resultEval.save()
            resultEval.heurPrincip.add(*heurPrincips)
            mergeScreenshots(resultEval,evals)
            # adding evaluators and evaluations id to the merged evaluation
            eval_ids = evals.values_list('pk' ,flat=True)
            evaluator_ids = evals.values_list('evaluator_id' , flat=True).distinct()
            resultEval.mergedFromEvaluations.add(*eval_ids)
            resultEval.merdedFromEvaluators.add(*evaluator_ids)

            # url = '/users/me/dashboard/project/update-merged-evaluation/' + str(resultEval.id) + '/'
            url = '/merge/project/' + str(resultEval.id) +'/update-merged-evaluation/'
            response = {'status' : 1 , 'message' : 'You can edit the result of merge!' , 'url' : url}
            return HttpResponse(json.dumps(response), content_type='application/json')
    return HttpResponse('not success')

def mergeFuncion( resultEval , evals):
    fields, heurPrincips = mergeFields(evals)
    resultEval.__dict__.update(fields)
    resultEval.save()
    resultEval.heurPrincip.add(*heurPrincips)
    mergeScreenshots(resultEval, evals)
    eval_ids = evals.values_list('pk', flat=True)
    evaluator_ids = evals.values_list('evaluator_id', flat=True).distinct()
    resultEval.mergedFromEvaluations.add(*eval_ids)
    resultEval.merdedFromEvaluators.add(*evaluator_ids)
    return resultEval


class UpdateMergedEvaluation(UpdateView):

    # eval = Evaluation.objects.get(pk=eval_id)
    # project = eval.ofProject
    template_name ='HE3/merged_evaluation_form.html'

    # form =MergeEvaluationForm(eval)
    model = Evaluation
    # fields = ['title','place', 'heurPrincip', 'description', 'recommendation', 'positivity' , 'severity' ,'frequency' ]
    form_class = MergeEvaluationForm

    def get_context_data(self, **kwargs):
        context = super(UpdateMergedEvaluation, self).get_context_data(**kwargs)
        eval = Evaluation.objects.get(pk=self.kwargs['pk'])
        context['project'] = eval.ofProject
        return context

    def get_form_kwargs(self):
        kwargs = super(UpdateMergedEvaluation, self).get_form_kwargs()
        kwargs['eval_id'] = self.kwargs['pk']
        return kwargs

    # return render(request, context={'project': project,'form' : form} ,template_name=template_name)

@login_required
def removeSelectedEval(request , eval_id ,se_eval_id):
    allSelected = get_object_or_404(Evaluation,pk=eval_id).mergedFromEvaluations
    if allSelected.count() > 1:
        allSelected.remove(get_object_or_404(Evaluation,pk=se_eval_id))
        messages.success(request , '<h4>Selected Evaluation is removed! </h4>')
    else:
        messages.error(request , '<h4>Selected Evaluation can not be removed. At least one evaluation should be in the merged evaluation! </h4>')

    return redirect(request.META.get('HTTP_REFERER'))

def makeNewMergeEval(request , eval_id , se_eval_id):
    mergedEval = get_object_or_404(Evaluation,pk=eval_id)
    allSelected = mergedEval.mergedFromEvaluations

    if allSelected.count() > 1 :
        mergedEval.mergedFromEvaluations.remove(get_object_or_404(Evaluation, pk=se_eval_id))
        resultEval = Evaluation(ofProject=mergedEval.ofProject, evaluator=mergedEval.ofProject.manager, merged=True)
        result = mergeFuncion(resultEval , allSelected.all())
        mergedEval.delete()
        messages.success(request , '<h4>  The selected evaluation is removed and a new merged Evaluation is created with the remaining selected evaluations! </h4>')
        return redirect('/merge/project/'+str(result.id)+'/update-merged-evaluation/')

    else:
        messages.error(request , '<h4>Selected Evaluation can not be removed. At least one evaluation should be in the merged evaluation! </h4>')
    return redirect(request.META.get('HTTP_REFERER'))



# ------------Report ----------------
def reportHtml(request , project_id):
    project = Project.objects.get(pk=project_id)
    mergedEvals = project.evaluation_for_project.filter(merged = True)
    evals = project.evaluation_for_project.filter(merged = False)
    evaluator_ids = evals.values_list('evaluator_id' , flat=True).distinct()
    evaluators = project.evaluators.filter(pk__in = evaluator_ids)
    environments = Environment.objects.filter(creator__in = evaluators).order_by('creator')

    context = { 'project' : project ,
                'evaluations' : evals.order_by('positivity' , 'evaluator' , 'severity'),
                'mergedEvals' : mergedEvals,
                'environments' :environments,
                'evaluators' : evaluators ,
                'pos_merge' :mergedEvals.filter(positivity = 'p'),
                'neg_merge' :mergedEvals.filter(positivity = 'n'),
    }

    return render (request , template_name='merge/report.html' , context=context)


# --------Export methods  ---------------------------------------------------------------

@login_required
def exportHtmlEvals(request, project_id, merge):
    project = Project.objects.get(pk=project_id)
    evals = project.evaluation_for_project.all()
    merged = False
    user = request.user
    if project.manager == user:
        if int(merge[0]) == 0:
            evals = evals.filter(merged=False)
        elif int(merge[0]) == 1:
            merged = True
            evals = evals.filter(merged=True)
    elif user in project.evaluators.all():
            evals = evals.filter(evaluator=request.user)
    else:
        return HttpResponse('<h2> Page not found!</h2>')


    template_name = 'merge/htmlexport.html'
    context = { 'evaluations' : evals ,
                'project': project ,
                'merged' :merged ,}

    return render(request,template_name=template_name,context=context)

# @login_required
# def exportDocFile(request, project_id , merge):
#     response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#     response['Content-Disposition'] = 'attachment; filename=demo.docx'
#     # merged = False
#     project = Project.objects.get(pk=project_id)
#     evals = project.evaluation_for_project.all()
#
#     if int(merge[0]) == 0:
#         evals = evals.filter(merged=False)
#     elif int(merge[0]) == 1:
#         # merged = True
#         evals = evals.filter(merged=True)
#
#     doc = fillDocFileEval(evals)
#     doc.save(response)
#
#     return response

# @login_required
# def exportCsvFile(request,list_id):
#     # # listofeval = ListOfEval.objects.get(pk=list_id)
#     # # evaluations = listofeval.evaluations.all().order_by('heurPrincip')
#     # # project = listofeval.ofProject
#     # response = HttpResponse(content_type='text/csv')
#     # response['Content-Disposition'] = 'attachment;filename = "Report.csv"'
#     #
#     # writer = csv.writer(response)
#     # writer.writerow([ 'Project: ', project.name,' Description: ' ,project.description,' link: ', project.link , ' Manager: ', project.manager])
#     # for e in evaluations:
#     #     writer.writerow([e.title ,
#     #                      e.place,
#     #                      e.heurPrincip,
#     #                      e.description,
#     #                      e.recommendation,
#     #                      e.get_positivity_display(),
#     #                      e.get_severity_display(),
#     #                      e.get_frequency_display(),
#     #                      e.tags,
#     #                      e.evaluator])
#     pass
#
#     # return response

# def fillDocFile(project):
#     doc = Document()
#     mergedEvals = project.evaluation_for_project.filter(merged=True)
#     evals = project.evaluation_for_project.filter(merged=False)
#     environments = Environment.objects.filter(creator__in=project.evaluators.all()).order_by('creator')
#     # context = {'project': project,
#     #            'evaluations': evals.order_by('positivity', 'evaluator', 'severity'),
#     #            'mergedEvals': mergedEvals,
#     #            'environments': environments,
#     #            'pos_merge': mergedEvals.filter(positivity='p'),
#     #            'neg_merge': mergedEvals.filter(positivity='n'),
#     #            }
#
#     doc.add_heading(str(project.name + 'Report'))
#     p = doc.add_paragraph(project.description)
#     p.add_run('bold').bold = True
#     p.add_run(' and some')
#     p.add_run('italic.').italic = True
#
#     doc.add_heading('Heading, level 1', level=1)
#     doc.add_paragraph('Intense quote', style='IntenseQuote')
#
#     doc.add_paragraph(
#         'first item in unordered list', style='ListBullet'
#     )
#     doc.add_paragraph(
#         'first item in ordered list', style='ListNumber'
#     )
#
#     table = doc.add_table(rows=1, cols=3)
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'Qty'
#     hdr_cells[1].text = 'Id'
#     hdr_cells[2].text = 'Desc'
#     # for item in listofeval.objects.all():
#     #     row_cells = table.add_row().cells
#     #     row_cells[0].text = str(item.name)
#     #     row_cells[1].text = str(item.genre)
#     #     row_cells[2].text = item.artist
#
#     doc.add_page_break()
#
#     return doc

# def fillDocFileEval(evals):
#     doc = Document()
#     doc.add_heading('Evaluations')
#
#     # table = doc.add_table(rows=evals.count(), cols=12)
#     for e in evals :
#         table = doc.add_table(rows=1, cols=11)
#
#         hdr_cells = table.rows[0].cells
#         hdr_cells[0].text = '#'
#         hdr_cells[1].text = 'Title'
#         hdr_cells[2].text = 'Heuristic Principle'
#         hdr_cells[3].text = 'Place'
#         hdr_cells[4].text = 'Link'
#         hdr_cells[7].text = 'Positivity'
#         hdr_cells[8].text = 'Severity'
#         hdr_cells[9].text = 'Frequency'
#         hdr_cells[10].text = 'Found By'
#
#     doc.add_heading('Description', level=1)
#     doc.add_heading('Recommendation', level=1)
#     doc.add_heading('Screenshot', level=1)
#
#
#     return doc

